"""DOCX export for chat transcripts, rendered with python-docx.

Like the PDF renderer this walks the shared block model from
``src/chat_export_model.py``, so the same transcript comes out with the same
structure in either format.

Design rules
------------
* **Real Word styles, not manual formatting.**  Headings use ``Heading 1-3``,
  lists use ``List Bullet`` / ``List Number`` (with the ``2``/``3`` variants for
  nesting), quotes use ``Quote``, tables use ``Table Grid``.  Everything else
  (role banners, code blocks, tool calls, metadata lines) gets a named
  ``Faustus ...`` style created once per document.  A reader can then restyle
  the whole export from Word's style pane instead of hand-editing paragraphs.
* **Hyperlinks are real hyperlinks.**  python-docx has no API for them, so we
  add the ``w:hyperlink`` element and its external relationship by hand
  (:func:`_add_hyperlink`) rather than dropping a blue-looking run that does
  not click.
* **Nothing raw reaches the XML.**  lxml escapes text for us, but control
  characters and lone surrogates are simply not representable in XML 1.0 and
  would raise, so they are stripped up front (:func:`_clean`).

Unlike the PDF, unicode needs no special handling here: a .docx stores UTF-8
text and Word does its own font fallback, so emoji and CJK survive intact.
"""

from __future__ import annotations

import io
import logging
import re
import unicodedata
from datetime import datetime
from types import SimpleNamespace
from typing import Any, Dict, List, Optional, Sequence

from src.chat_export import is_document
from src.chat_export_model import (
    Block,
    ExportMessage,
    ExportUnavailable,
    Span,
    ToolCall,
    Transcript,
)

logger = logging.getLogger(__name__)

DOCX_MISSING = (
    "DOCX export requires the 'python-docx' package. Install it with "
    "`pip install python-docx` (or `pip install -r requirements.txt`)."
)

# Kept in step with src/chat_export_pdf.py so both formats read identically.
LABELS = {
    "user": "User",
    "assistant": "Assistant",
    "system": "System",
    "tool": "Tool",
    "model": "Model",
    "exported": "Exported",
    "messages": "messages",
    "message": "message",
    "session": "Session",
    "project": "Project",
    "workspace": "Workspace",
    "attachments": "Attachments",
    "empty": "(this conversation has no messages)",
    "tool_call": "tool call",
    "arguments": "arguments",
    "result": "result",
    "truncated": "... [truncated]",
    "image": "[image]",
    "page": "Page ",
    "page_of": " of ",
}

ROLE_COLORS = {
    "user": ("1D4ED8", "EFF6FF"),
    "assistant": ("047857", "ECFDF5"),
    "system": ("4B5563", "F3F4F6"),
    "tool": ("92400E", "FFFBEB"),
}

CODE_FILL = "F6F8FA"
TOOL_FILL = "FFFBEB"
MUTED = "6B7280"
RULE = "D1D5DB"

MONO_FONT = "Consolas"           # present on Windows; Word substitutes elsewhere
CODE_FONT_SIZE_PT = 8.5
TOOL_TEXT_LIMIT = 2000

SAFE_LINK_SCHEMES = ("http://", "https://", "mailto:", "ftp://", "ftps://", "tel:")

# Control characters and lone surrogates are not representable in XML 1.0;
# lxml raises ValueError on them, which would abort the export.
_ILLEGAL_XML = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f\ud800-\udfff]")

# Child order inside <w:pPr> is fixed by the schema, so new elements are
# inserted before their successors rather than appended. (Mirrors the private
# sequence in docx.oxml.text.parfmt.CT_PPr.)
_AFTER_PBDR = ("w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr")
_AFTER_SHD = ("w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr")
_AFTER_RUN_SHD = ("w:fitText", "w:vertAlign", "w:rtl", "w:cs", "w:em", "w:lang",
                  "w:eastAsianLayout", "w:specVanish", "w:oMath")


# ---------------------------------------------------------------------------
# lazy python-docx import
# ---------------------------------------------------------------------------

_DX: Optional[SimpleNamespace] = None


def _dx() -> SimpleNamespace:
    """Import python-docx on first use.

    Kept out of module scope so importing this module costs nothing when
    python-docx is absent, and so the caller sees an ``ExportUnavailable``
    naming the package instead of an ImportError from deep in the pipeline.
    """
    global _DX
    if _DX is not None:
        return _DX
    try:
        import docx
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.opc.constants import RELATIONSHIP_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover - exercised with a patched import
        raise ExportUnavailable(DOCX_MISSING) from exc

    _DX = SimpleNamespace(
        docx=docx,
        WD_STYLE_TYPE=WD_STYLE_TYPE,
        WD_TABLE_ALIGNMENT=WD_TABLE_ALIGNMENT,
        WD_ALIGN_PARAGRAPH=WD_ALIGN_PARAGRAPH,
        RELATIONSHIP_TYPE=RELATIONSHIP_TYPE,
        OxmlElement=OxmlElement,
        qn=qn,
        Cm=Cm,
        Pt=Pt,
        RGBColor=RGBColor,
    )
    return _DX


# ---------------------------------------------------------------------------
# small shared helpers (twins of the ones in src/chat_export_pdf.py; each
# renderer stays importable on its own, with no dependency on the other)
# ---------------------------------------------------------------------------


def _clean(text: Any) -> str:
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    return _ILLEGAL_XML.sub("", text)


def _safe_href(href: Any) -> str:
    url = _clean(href).strip()
    if not url:
        return ""
    return url if url.lower().startswith(SAFE_LINK_SCHEMES) else ""


def _format_timestamp(raw: Any) -> str:
    text = _clean(raw).strip()
    if not text:
        return ""
    try:
        return datetime.fromisoformat(text.replace("Z", "+00:00")).strftime("%Y-%m-%d %H:%M")
    except (ValueError, TypeError):
        return text[:40]


def _display_width(text: str) -> int:
    return sum(2 if unicodedata.east_asian_width(c) in ("W", "F") else 1 for c in text)


def _wrap_code(text: str, max_cols: int) -> str:
    """Hard-wrap code lines.

    Word will wrap an over-long run of characters on its own, but where it
    breaks depends on the reader's page setup and font substitution. Wrapping
    here keeps a minified line or a long URL identical to the PDF output.
    """
    max_cols = max(20, int(max_cols))
    out: List[str] = []
    for line in _clean(text).replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = line.replace("\t", "    ")
        if _display_width(line) <= max_cols:
            out.append(line)
            continue
        current: List[str] = []
        width = 0
        for ch in line:
            char_width = 2 if unicodedata.east_asian_width(ch) in ("W", "F") else 1
            if width + char_width > max_cols and current:
                out.append("".join(current))
                current, width = [], 0
            current.append(ch)
            width += char_width
        out.append("".join(current))
    return "\n".join(out)


def _truncate(text: str, limit: int) -> str:
    text = _clean(text)
    return text if len(text) <= limit else text[:limit] + "\n" + LABELS["truncated"]


def _spans_text(spans: Sequence[Span]) -> str:
    return "".join(_clean(getattr(s, "text", s)) for s in spans or ())


# ---------------------------------------------------------------------------
# raw XML helpers
# ---------------------------------------------------------------------------


def _shading(fill: str):
    dx = _dx()
    shd = dx.OxmlElement("w:shd")
    shd.set(dx.qn("w:val"), "clear")
    shd.set(dx.qn("w:color"), "auto")
    shd.set(dx.qn("w:fill"), fill)
    return shd


def _border(edge: str, color: str, size: int = 6, space: int = 4):
    dx = _dx()
    element = dx.OxmlElement("w:" + edge)
    element.set(dx.qn("w:val"), "single")
    element.set(dx.qn("w:sz"), str(size))
    element.set(dx.qn("w:space"), str(space))
    element.set(dx.qn("w:color"), color)
    return element


def _style_shading(style, fill: str) -> None:
    pPr = style.element.get_or_add_pPr()
    pPr.insert_element_before(_shading(fill), *_AFTER_SHD)


def _style_border(style, edges: Dict[str, str]) -> None:
    dx = _dx()
    pPr = style.element.get_or_add_pPr()
    pBdr = dx.OxmlElement("w:pBdr")
    for edge, color in edges.items():
        pBdr.append(_border(edge, color))
    pPr.insert_element_before(pBdr, *_AFTER_PBDR)


def _run_font(style, name: str) -> None:
    """Set an explicit font on a style, including the complex-script slot."""
    dx = _dx()
    style.font.name = name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(dx.qn("w:rFonts"))
    if rFonts is None:  # pragma: no cover - python-docx creates it via font.name
        rFonts = dx.OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(dx.qn("w:cs"), name)


def _add_hyperlink(paragraph, url: str, runs: Sequence[Any]) -> None:
    """Wrap already-created runs in a real ``w:hyperlink`` element.

    python-docx exposes no hyperlink API: the relationship has to be added to
    the document part and referenced by id from the paragraph. Building the
    runs first and then re-parenting their XML keeps python-docx's own text
    handling (breaks, xml:space) intact.
    """
    dx = _dx()
    rel_id = paragraph.part.relate_to(url, dx.RELATIONSHIP_TYPE.HYPERLINK,
                                      is_external=True)
    link = dx.OxmlElement("w:hyperlink")
    link.set(dx.qn("r:id"), rel_id)
    for run in runs:
        link.append(run._r)
    paragraph._p.append(link)


def _field(paragraph, instruction: str, placeholder: str) -> None:
    """Append a simple Word field (used for PAGE / NUMPAGES in the footer)."""
    dx = _dx()
    field = dx.OxmlElement("w:fldSimple")
    field.set(dx.qn("w:instr"), instruction)
    run = dx.OxmlElement("w:r")
    text = dx.OxmlElement("w:t")
    text.text = placeholder
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def _repeat_header_row(row) -> None:
    dx = _dx()
    trPr = row._tr.get_or_add_trPr()
    trPr.append(dx.OxmlElement("w:tblHeader"))


# ---------------------------------------------------------------------------
# document styles
# ---------------------------------------------------------------------------


def _ensure_style(document, name: str, style_type, base: Optional[str] = None):
    """Fetch a style by name, creating it when the template lacks it."""
    try:
        return document.styles[name]
    except KeyError:
        pass
    style = document.styles.add_style(name, style_type)
    if base:
        try:
            style.base_style = document.styles[base]
        except KeyError:  # pragma: no cover - defensive
            pass
    return style


def _build_styles(document) -> Dict[str, Any]:
    """Create the Faustus styles once, so the body can stay style-driven."""
    dx = _dx()
    paragraph = dx.WD_STYLE_TYPE.PARAGRAPH
    character = dx.WD_STYLE_TYPE.CHARACTER
    styles: Dict[str, Any] = {}

    meta = _ensure_style(document, "Faustus Meta", paragraph, "Normal")
    meta.font.size = dx.Pt(8.5)
    meta.font.color.rgb = dx.RGBColor.from_string(MUTED)
    meta.paragraph_format.space_after = dx.Pt(2)
    styles["meta"] = meta

    for role, (color, fill) in ROLE_COLORS.items():
        name = "Faustus Role %s" % role.capitalize()
        style = _ensure_style(document, name, paragraph, "Normal")
        style.font.bold = True
        style.font.size = dx.Pt(10)
        style.font.color.rgb = dx.RGBColor.from_string(color)
        style.paragraph_format.space_before = dx.Pt(14)
        style.paragraph_format.space_after = dx.Pt(4)
        style.paragraph_format.keep_with_next = True
        _style_shading(style, fill)
        styles["role_" + role] = style

    code = _ensure_style(document, "Faustus Code", paragraph, "Normal")
    _run_font(code, MONO_FONT)
    code.font.size = dx.Pt(CODE_FONT_SIZE_PT)
    code.paragraph_format.space_before = dx.Pt(4)
    code.paragraph_format.space_after = dx.Pt(8)
    code.paragraph_format.left_indent = dx.Cm(0.3)
    _style_shading(code, CODE_FILL)
    styles["code"] = code

    tool = _ensure_style(document, "Faustus Tool Call", paragraph, "Normal")
    _run_font(tool, MONO_FONT)
    tool.font.size = dx.Pt(8)
    tool.font.color.rgb = dx.RGBColor.from_string("5B3A0B")
    tool.paragraph_format.space_before = dx.Pt(4)
    tool.paragraph_format.space_after = dx.Pt(6)
    tool.paragraph_format.left_indent = dx.Cm(0.3)
    _style_shading(tool, TOOL_FILL)
    styles["tool"] = tool

    rule = _ensure_style(document, "Faustus Rule", paragraph, "Normal")
    rule.font.size = dx.Pt(1)
    rule.paragraph_format.space_before = dx.Pt(6)
    rule.paragraph_format.space_after = dx.Pt(6)
    _style_border(rule, {"bottom": RULE})
    styles["rule"] = rule

    inline = _ensure_style(document, "Faustus Inline Code", character)
    _run_font(inline, MONO_FONT)
    inline.font.size = dx.Pt(9)
    styles["inline_code"] = inline

    # Word's own Hyperlink character style; absent from python-docx's default
    # template, so create it rather than leave links unstyled.
    link = _ensure_style(document, "Hyperlink", character)
    link.font.color.rgb = dx.RGBColor.from_string("0563C1")
    link.font.underline = True
    styles["hyperlink"] = link
    return styles


# ---------------------------------------------------------------------------
# inline spans
# ---------------------------------------------------------------------------


def _add_spans(paragraph, spans: Sequence[Span], ctx: "_Ctx") -> None:
    """Append inline spans to a paragraph, links included."""
    for span in spans or ():
        if not isinstance(span, Span):
            text = _clean(span)
            if text:
                paragraph.add_run(text)
            continue
        text = _clean(span.text)
        if not text:
            continue
        run = paragraph.add_run(text)
        if span.code:
            run.style = ctx.styles["inline_code"]
        if span.bold:
            run.bold = True
        if span.italic:
            run.italic = True
        if span.strike:
            run.font.strike = True
        href = _safe_href(span.href)
        if href:
            if not span.code:
                run.style = ctx.styles["hyperlink"]
            else:
                run.font.underline = True
                run.font.color.rgb = _dx().RGBColor.from_string("0563C1")
            _add_hyperlink(paragraph, href, [run])


# ---------------------------------------------------------------------------
# blocks
# ---------------------------------------------------------------------------


class _Ctx:
    """Per-export state: the document, its styles and the usable text width."""

    def __init__(self, document, styles: Dict[str, Any], text_width_pt: float,
                 document_mode: bool = False):
        self.document = document
        self.styles = styles
        self.text_width_pt = text_width_pt
        # A document has no speakers and no message count to report; see
        # ``DOCUMENT_FLAG`` in src/chat_export.py.
        self.document_mode = document_mode


def _add_paragraph(ctx: "_Ctx", style=None, indent_cm: Optional[float] = 0.0):
    """Add a paragraph. ``indent_cm=None`` keeps the style's own indentation.

    List and Quote styles carry their own (hanging) indents; overriding them
    with the message indent would flatten the list levels, so those callers
    pass None.
    """
    dx = _dx()
    paragraph = ctx.document.add_paragraph()
    if style is not None:
        paragraph.style = style
    if indent_cm:
        paragraph.paragraph_format.left_indent = dx.Cm(indent_cm)
    return paragraph


def _add_code(ctx: "_Ctx", text: str, lang: str, indent_cm: float) -> None:
    dx = _dx()
    lang = _clean(lang).strip()
    if lang:
        label = _add_paragraph(ctx, ctx.styles["meta"], indent_cm + 0.3)
        label.paragraph_format.space_after = dx.Pt(0)
        run = label.add_run(lang)
        run.font.size = dx.Pt(7.5)
    # ~0.62 em per monospaced character is a good approximation across
    # Consolas/Courier/DejaVu Sans Mono; the wrap only needs to be close.
    columns = int((ctx.text_width_pt - indent_cm * 28.35) / (CODE_FONT_SIZE_PT * 0.62))
    paragraph = _add_paragraph(ctx, ctx.styles["code"], indent_cm)
    # python-docx turns "\n" into <w:br/> and marks the runs xml:space="preserve",
    # so indentation and blank lines inside the block survive.
    paragraph.add_run(_wrap_code(text, columns))


def _add_table(ctx: "_Ctx", block: Block, indent_cm: float) -> None:
    dx = _dx()
    rows = [row for row in (block.rows or []) if row is not None]
    if not rows:
        return
    ncols = max((len(row) for row in rows), default=0)
    if ncols <= 0:
        return
    table = ctx.document.add_table(rows=len(rows), cols=ncols)
    try:
        table.style = ctx.document.styles["Table Grid"]
    except KeyError:  # pragma: no cover - present in the default template
        pass
    table.alignment = dx.WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for index, row in enumerate(rows):
        for col in range(ncols):
            cell = table.cell(index, col)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = dx.Pt(2)
            spans = row[col] if col < len(row) else []
            _add_spans(paragraph, spans, ctx)
            if block.header and index == 0:
                for run in paragraph.runs:
                    run.bold = True
    if block.header and rows:
        _repeat_header_row(table.rows[0])
    # A table swallows the following paragraph's spacing; add breathing room.
    spacer = _add_paragraph(ctx, ctx.styles["meta"])
    spacer.paragraph_format.space_after = dx.Pt(2)


def _list_style_name(ordered: bool, level: int) -> str:
    base = "List Number" if ordered else "List Bullet"
    level = max(1, min(3, level))
    return base if level == 1 else "%s %d" % (base, level)


def _add_blocks(ctx: "_Ctx", blocks: Sequence[Block], indent_cm: float = 0.0,
                level: int = 1, quote: bool = False) -> None:
    for block in blocks or ():
        try:
            _add_block(ctx, block, indent_cm, level, quote)
        except Exception:
            # One bad block must not cost the user the whole export.
            logger.exception("chat_export_docx: skipping unrenderable block %r",
                             getattr(block, "kind", "?"))


def _add_block(ctx: "_Ctx", block: Block, indent_cm: float, level: int,
               quote: bool) -> None:
    kind = getattr(block, "kind", "") or "para"

    if kind == "heading":
        heading_level = block.level if isinstance(block.level, int) else 1
        heading_level = min(3, max(1, heading_level))
        paragraph = _add_paragraph(ctx, "Heading %d" % heading_level, indent_cm)
        _add_spans(paragraph, block.spans, ctx)
        return

    if kind == "code":
        _add_code(ctx, block.text, block.lang, indent_cm)
        return

    if kind == "list":
        for item in block.items or []:
            first = True
            for child in item or []:
                if first and getattr(child, "kind", "para") == "para":
                    paragraph = _add_paragraph(
                        ctx, _list_style_name(block.ordered, level), None)
                    _add_spans(paragraph, child.spans, ctx)
                    first = False
                else:
                    _add_blocks(ctx, [child], indent_cm + 0.6, level + 1, quote)
            if first:  # an item with no leading paragraph
                _add_paragraph(ctx, _list_style_name(block.ordered, level), None)
        return

    if kind == "quote":
        _add_blocks(ctx, block.children, indent_cm, level, quote=True)
        return

    if kind == "table":
        _add_table(ctx, block, indent_cm)
        return

    if kind == "hr":
        _add_paragraph(ctx, ctx.styles["rule"], indent_cm)
        return

    if kind == "image":
        # Images are not embedded: an export must not fetch remote URLs, and a
        # local path may be gone by the time the document is opened.
        alt = _spans_text(block.spans) or _clean(block.text)
        paragraph = _add_paragraph(ctx, ctx.styles["meta"], indent_cm)
        label = ("%s %s" % (LABELS["image"], alt)).strip()
        run = paragraph.add_run(label)
        run.italic = True
        href = _safe_href(block.href)
        if href:
            link_run = paragraph.add_run("  " + href)
            link_run.style = ctx.styles["hyperlink"]
            _add_hyperlink(paragraph, href, [link_run])
        return

    # "para" and anything unknown.
    spans = block.spans
    if not spans and block.text:
        spans = [Span(block.text)]
    if not spans:
        return
    paragraph = _add_paragraph(ctx, "Quote" if quote else None,
                               None if quote else indent_cm)
    _add_spans(paragraph, spans, ctx)


# ---------------------------------------------------------------------------
# messages
# ---------------------------------------------------------------------------


def _role_key(role: Any) -> str:
    role = _clean(role).strip().lower()
    return role if role in ROLE_COLORS else "system"


def _role_label(role: Any) -> str:
    key = _clean(role).strip().lower()
    if key in LABELS:
        return LABELS[key]
    return key.replace("_", " ").title() if key else LABELS["system"]


def _add_tool_call(ctx: "_Ctx", call: ToolCall, indent_cm: float) -> None:
    name = _clean(getattr(call, "name", "")) or "?"
    head = "%s %s" % (LABELS["tool_call"], name)
    bits = []
    status = _clean(getattr(call, "status", ""))
    if status:
        bits.append(status)
    duration = getattr(call, "duration_s", None)
    if isinstance(duration, (int, float)):
        bits.append("%.2fs" % duration)
    if bits:
        head += "  (%s)" % ", ".join(bits)

    lines = [head]
    arguments = _truncate(getattr(call, "arguments", "") or "", TOOL_TEXT_LIMIT).strip()
    if arguments:
        lines.append("%s: %s" % (LABELS["arguments"], arguments))
    result = _truncate(getattr(call, "result", "") or "", TOOL_TEXT_LIMIT).strip()
    if result:
        lines.append("%s: %s" % (LABELS["result"], result))

    columns = int((ctx.text_width_pt - indent_cm * 28.35) / (8 * 0.62))
    paragraph = _add_paragraph(ctx, ctx.styles["tool"], indent_cm)
    paragraph.add_run(_wrap_code("\n".join(lines), columns))


def _add_message(ctx: "_Ctx", message: ExportMessage) -> None:
    indent_cm = 0.35
    if not ctx.document_mode:
        role = _role_key(getattr(message, "role", ""))
        banner = _add_paragraph(ctx, ctx.styles["role_" + role])
        banner.add_run(_role_label(getattr(message, "role", "")))

        trailing = []
        timestamp = _format_timestamp(getattr(message, "timestamp", ""))
        if timestamp:
            trailing.append(timestamp)
        model = _clean(getattr(message, "model", "")).strip()
        if model:
            trailing.append(model)
        if trailing:
            run = banner.add_run("  ·  " + "  ·  ".join(trailing))
            run.bold = False
            run.font.size = _dx().Pt(8.5)
            run.font.color.rgb = _dx().RGBColor.from_string(MUTED)

    _add_blocks(ctx, getattr(message, "blocks", None) or [], indent_cm)

    for call in getattr(message, "tool_calls", None) or []:
        try:
            _add_tool_call(ctx, call, indent_cm)
        except Exception:
            logger.exception("chat_export_docx: skipping unrenderable tool call")

    attachments = [_clean(a) for a in (getattr(message, "attachments", None) or []) if _clean(a)]
    if attachments:
        paragraph = _add_paragraph(ctx, ctx.styles["meta"], indent_cm)
        run = paragraph.add_run("%s: %s" % (LABELS["attachments"], ", ".join(attachments)))
        run.italic = True


def _add_header(ctx: "_Ctx", transcript: Transcript) -> None:
    name = _clean(getattr(transcript, "name", "")).strip() or "Conversation"
    title = _add_paragraph(ctx, "Title")
    title.add_run(name)

    messages = list(getattr(transcript, "messages", None) or [])
    meta: List[str] = []
    model = _clean(getattr(transcript, "model", "")).strip()
    if model:
        meta.append("%s: %s" % (LABELS["model"], model))
    # "1 message" is true of a document only as an accident of how it travels.
    if not ctx.document_mode:
        meta.append("%d %s" % (len(messages),
                               LABELS["messages"] if len(messages) != 1 else LABELS["message"]))
    exported_at = getattr(transcript, "exported_at", None)
    if isinstance(exported_at, datetime):
        meta.append("%s %s" % (LABELS["exported"], exported_at.strftime("%Y-%m-%d %H:%M")))
    elif exported_at:
        meta.append("%s %s" % (LABELS["exported"], _clean(exported_at)))
    for key in ("project", "workspace", "session_id"):
        value = _clean(getattr(transcript, key, "")).strip()
        if value:
            meta.append("%s: %s" % (LABELS["session" if key == "session_id" else key], value))

    # Only reachable for a document: a conversation always has its count.
    if meta:
        paragraph = _add_paragraph(ctx, ctx.styles["meta"])
        paragraph.add_run("  ·  ".join(meta))
    _add_paragraph(ctx, ctx.styles["rule"])


def _setup_page(document, name: str) -> float:
    """A4 with 2 cm margins, a footer with page fields. Returns text width in pt."""
    dx = _dx()
    section = document.sections[0]
    section.page_width = dx.Cm(21)
    section.page_height = dx.Cm(29.7)
    for attribute in ("left_margin", "right_margin"):
        setattr(section, attribute, dx.Cm(2))
    section.top_margin = dx.Cm(1.8)
    section.bottom_margin = dx.Cm(1.8)

    footer = section.footer.paragraphs[0]
    footer.style = document.styles["Faustus Meta"]
    footer.alignment = dx.WD_ALIGN_PARAGRAPH.RIGHT
    label = name if len(name) <= 70 else name[:69] + "…"
    footer.add_run(label + "    " + LABELS["page"])
    _field(footer, " PAGE ", "1")
    footer.add_run(LABELS["page_of"])
    _field(footer, " NUMPAGES ", "1")

    core = document.core_properties
    core.title = name
    core.author = "Faustus"
    core.comments = "Chat transcript exported by Faustus"
    return float(section.page_width - section.left_margin - section.right_margin) / 12700.0


# ---------------------------------------------------------------------------
# public API
# ---------------------------------------------------------------------------


def render(transcript: Transcript) -> bytes:
    """Render a transcript as a .docx document.

    Raises :class:`ExportUnavailable` (never a bare ImportError) when
    python-docx is not installed.
    """
    dx = _dx()
    document = dx.docx.Document()
    styles = _build_styles(document)
    name = _clean(getattr(transcript, "name", "")).strip() or "Conversation"
    text_width_pt = _setup_page(document, name)
    ctx = _Ctx(document, styles, text_width_pt, document_mode=is_document(transcript))

    _add_header(ctx, transcript)
    messages = list(getattr(transcript, "messages", None) or [])
    if not messages:
        paragraph = _add_paragraph(ctx, styles["meta"])
        paragraph.add_run(LABELS["empty"]).italic = True
    for message in messages:
        _add_message(ctx, message)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
